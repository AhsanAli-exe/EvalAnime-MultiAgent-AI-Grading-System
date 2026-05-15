import zipfile
from pathlib import Path
from pypdf import PdfReader
from docx import Document
from ..storage import append_event

def _ext(path):
    return Path(path).suffix.lower().lstrip(".")

def parse_pdf(run_id,file_path):
    append_event(run_id,"tool_call_start",{"tool":"parse_pdf","file":file_path})
    text_parts=[]
    page_count=0
    try:
        reader=PdfReader(file_path)
        page_count=len(reader.pages)
        for page in reader.pages:
            t=page.extract_text() or ""
            text_parts.append(t)
        text="\n".join(text_parts).strip()
        result={"ok":True,"text":text,"page_count":page_count,"sparse":len(text)<40}
    except Exception as e:
        result={"ok":False,"error":str(e),"text":"","page_count":0,"sparse":True}
    append_event(run_id,"tool_call_end",{"tool":"parse_pdf","ok":result["ok"],"chars":len(result["text"]),"pages":result["page_count"],"sparse":result["sparse"]})
    return result

def parse_docx(run_id,file_path):
    append_event(run_id,"tool_call_start",{"tool":"parse_docx","file":file_path})
    try:
        doc=Document(file_path)
        lines=[p.text for p in doc.paragraphs]
        for t in doc.tables:
            for row in t.rows:
                for cell in row.cells:
                    if cell.text and cell.text.strip():
                        lines.append(cell.text)
        text="\n".join(lines).strip()
        result={"ok":True,"text":text,"paragraphs":len(doc.paragraphs),"tables":len(doc.tables)}
    except Exception as e:
        result={"ok":False,"error":str(e),"text":"","paragraphs":0,"tables":0}
    append_event(run_id,"tool_call_end",{"tool":"parse_docx","ok":result["ok"],"chars":len(result["text"])})
    return result

def unzip_archive(run_id,file_path,extract_to=None):
    append_event(run_id,"tool_call_start",{"tool":"unzip_archive","file":file_path})
    out_dir=Path(extract_to) if extract_to else Path(file_path).with_suffix("")
    out_dir.mkdir(parents=True,exist_ok=True)
    inner_files=[]
    try:
        with zipfile.ZipFile(file_path,"r") as z:
            z.extractall(out_dir)
            inner_files=[str(out_dir/name) for name in z.namelist() if not name.endswith("/")]
        result={"ok":True,"extracted_to":str(out_dir),"files":inner_files}
    except Exception as e:
        result={"ok":False,"error":str(e),"extracted_to":str(out_dir),"files":[]}
    append_event(run_id,"tool_call_end",{"tool":"unzip_archive","ok":result["ok"],"files":inner_files})
    return result

def parse_any(run_id,file_path):
    e=_ext(file_path)
    if e=="pdf":
        return {"kind":"pdf","result":parse_pdf(run_id,file_path)}
    if e=="docx":
        return {"kind":"docx","result":parse_docx(run_id,file_path)}
    if e=="zip":
        info=unzip_archive(run_id,file_path)
        inner_pdf=next((f for f in info["files"] if f.lower().endswith(".pdf")),None)
        inner_docx=next((f for f in info["files"] if f.lower().endswith(".docx")),None)
        if inner_pdf:
            r=parse_pdf(run_id,inner_pdf)
            return {"kind":"zip+pdf","inner":inner_pdf,"result":r}
        if inner_docx:
            r=parse_docx(run_id,inner_docx)
            return {"kind":"zip+docx","inner":inner_docx,"result":r}
        return {"kind":"zip","result":{"ok":False,"error":"no supported file inside zip","text":""}}
    return {"kind":"unknown","result":{"ok":False,"error":f"unsupported extension: {e}","text":""}}


# helper notes (read-me at a glance):
# _ext(path)          -> just gets the file extension in lowercase (no dot)
# parse_pdf()         -> opens a pdf and pulls out plain text. if the text is shorter
#                        than 40 chars we mark sparse=True so the agent knows to fall
#                        back to OCR (image-only pdfs land here, e.g. student #5).
# parse_docx()        -> reads a .docx file paragraph by paragraph into one string
# unzip_archive()     -> extracts a zip next to the file and lists what came out
# parse_any()         -> the "smart" entry point. given any file, pick the right parser
#                        automatically. for a zip it also peeks inside to grab the first
#                        pdf or docx it finds. every step logs tool_call_start /
#                        tool_call_end events so the audit log shows exactly what ran.
